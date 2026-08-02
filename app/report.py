from __future__ import annotations

from datetime import datetime
from pathlib import Path
from typing import Any

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

from .storage import MetadataStore


def pct(value: Any) -> str:
    if value is None:
        return "—"
    return f"{float(value) * 100:.1f}%"


def text(value: Any) -> str:
    if value is None or value == "":
        return "—"
    return str(value)


def generate_docx_report(store: MetadataStore, snapshot_id: str, report_dir: Path) -> Path:
    snapshot_rows = store.query(
        """
        SELECT ss.*, ds.name AS source_name, ds.dialect
        FROM scan_snapshot ss
        JOIN data_source ds ON ds.source_id = ss.source_id
        WHERE ss.snapshot_id = ?
        """,
        [snapshot_id],
    )
    if not snapshot_rows:
        raise ValueError("快照不存在。")
    snapshot = snapshot_rows[0]
    overview = store.query(
        """
        SELECT
          (SELECT COUNT(*) FROM meta_table WHERE snapshot_id = ?) AS table_count,
          (SELECT COALESCE(SUM(row_count), 0) FROM table_stat WHERE snapshot_id = ?) AS total_rows,
          (SELECT COUNT(*) FROM meta_column WHERE snapshot_id = ?) AS column_count,
          (SELECT AVG(avg_fill_rate) FROM table_stat WHERE snapshot_id = ?) AS avg_fill_rate,
          (SELECT AVG(avg_valid_rate) FROM table_stat WHERE snapshot_id = ?) AS avg_valid_rate,
          (SELECT COUNT(*) FROM meta_column WHERE snapshot_id = ? AND is_sensitive) AS sensitive_columns
        """,
        [snapshot_id, snapshot_id, snapshot_id, snapshot_id, snapshot_id, snapshot_id],
    )[0]
    low_columns = store.query(
        """
        SELECT mc.table_name, mc.column_name, mc.data_type, cs.fill_rate, cs.valid_rate, cs.skipped_reason
        FROM column_stat cs
        JOIN meta_column mc ON mc.snapshot_id = cs.snapshot_id AND mc.column_id = cs.column_id
        WHERE cs.snapshot_id = ? AND cs.fill_rate IS NOT NULL
        ORDER BY cs.fill_rate ASC
        LIMIT 20
        """,
        [snapshot_id],
    )
    gap_columns = store.query(
        """
        SELECT mc.table_name, mc.column_name, cs.fill_rate, cs.valid_rate,
               (cs.fill_rate - cs.valid_rate) AS gap
        FROM column_stat cs
        JOIN meta_column mc ON mc.snapshot_id = cs.snapshot_id AND mc.column_id = cs.column_id
        WHERE cs.snapshot_id = ? AND cs.fill_rate IS NOT NULL AND cs.valid_rate IS NOT NULL
        ORDER BY gap DESC
        LIMIT 20
        """,
        [snapshot_id],
    )
    tables = store.query(
        """
        SELECT mt.table_name, mt.column_count, ts.row_count, ts.avg_fill_rate, ts.avg_valid_rate,
               ts.date_column, ts.min_date, ts.max_date
        FROM meta_table mt
        LEFT JOIN table_stat ts ON ts.snapshot_id = mt.snapshot_id AND ts.table_id = mt.table_id
        WHERE mt.snapshot_id = ?
        ORDER BY ts.row_count DESC NULLS LAST, mt.table_name
        """,
        [snapshot_id],
    )
    metrics = store.query(
        """
        SELECT metric_code, name, definition, formula, denominator, boundary
        FROM metric_registry
        WHERE version = ?
        ORDER BY metric_code
        """,
        [snapshot["metric_def_version"]],
    )

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    title = doc.add_heading("DataPulse 数据质量报告初稿", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"数据源：{snapshot['source_name']} ｜ 方言：{snapshot['dialect']} ｜ 快照：{snapshot_id}").bold = True
    doc.add_paragraph(
        f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}；"
        f"口径版本：{snapshot['metric_def_version']}；"
        f"快照状态：{snapshot['status']}。"
    )

    doc.add_heading("1. 口径声明", level=1)
    doc.add_paragraph("本报告不输出综合质量评分，只输出分项指标与问题榜单。所有指标均来自本地 DuckDB 元数据仓库，报告生成过程不连接源库。")

    doc.add_heading("2. 库概览", level=1)
    overview_table = doc.add_table(rows=1, cols=2)
    overview_table.style = "Table Grid"
    overview_table.rows[0].cells[0].text = "指标"
    overview_table.rows[0].cells[1].text = "值"
    for label, value in [
        ("表数量", overview["table_count"]),
        ("总行数", overview["total_rows"]),
        ("字段数", overview["column_count"]),
        ("平均有值率", pct(overview["avg_fill_rate"])),
        ("平均有效率", pct(overview["avg_valid_rate"])),
        ("敏感字段数", overview["sensitive_columns"]),
    ]:
        row = overview_table.add_row().cells
        row[0].text = label
        row[1].text = text(value)

    doc.add_heading("3. 问题榜单", level=1)
    doc.add_paragraph("有值率最低字段")
    add_table(doc, ["表", "字段", "类型", "有值率", "有效率", "跳过原因"], [
        [r["table_name"], r["column_name"], r["data_type"], pct(r["fill_rate"]), pct(r["valid_rate"]), text(r["skipped_reason"])]
        for r in low_columns
    ])
    doc.add_paragraph("有值率 / 有效率落差榜")
    add_table(doc, ["表", "字段", "有值率", "有效率", "落差"], [
        [r["table_name"], r["column_name"], pct(r["fill_rate"]), pct(r["valid_rate"]), pct(r["gap"])]
        for r in gap_columns
    ])

    doc.add_heading("4. 表级明细附录", level=1)
    add_table(doc, ["表", "行数", "字段数", "平均有值率", "平均有效率", "业务日期字段", "日期范围"], [
        [
            r["table_name"],
            text(r["row_count"]),
            text(r["column_count"]),
            pct(r["avg_fill_rate"]),
            pct(r["avg_valid_rate"]),
            text(r["date_column"]),
            f"{text(r['min_date'])} ~ {text(r['max_date'])}",
        ]
        for r in tables
    ])

    doc.add_heading("5. 指标口径附录", level=1)
    add_table(doc, ["编码", "名称", "定义", "公式", "分母", "边界"], [
        [m["metric_code"], m["name"], m["definition"], m["formula"], m["denominator"], m["boundary"]]
        for m in metrics
    ])

    report_dir.mkdir(parents=True, exist_ok=True)
    path = report_dir / f"DataPulse_report_{snapshot_id[:8]}.docx"
    doc.save(path)
    store.add_audit("report_exported", {"snapshot_id": snapshot_id, "path": str(path)})
    return path


def add_table(doc: Document, headers: list[str], rows: list[list[Any]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    if not rows:
        cells = table.add_row().cells
        cells[0].text = "无数据"
        return
    for item in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(item):
            cells[idx].text = text(value)
